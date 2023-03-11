from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from random import randint, choice
from pprint import pprint

from docx2pdf import convert

def system_translation(number, radixresult):
    digits = "0123456789ABCDEF"
    res = ''
    number = abs(number)
    while number != 0:
        res = str(digits[number % radixresult]) + res
        number //= radixresult
    return res


def generate_exercise(numbertasks, operation, radix):
    res = [(system_translation(randint(20, 512), radix),
           system_translation(randint(20, 512), radix)) for i in range(numbertasks)]
    resanswers = [system_translation(eval(str(int(item[0], radix)) + operation +
                       str(int(item[1], radix))), radix) for item in res]
    print()
    print()
    print(res)
    print(operation)
    print()
    print()
    return res, resanswers


def generate_pdf(theme, number_of_exercises, operations, radixs):

    answers_document = Document()
    answers_nextp = answers_document.add_paragraph()
    answers_nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answers_nextrun = answers_nextp.add_run("Ответы")
    answers_nextrun.font.size = Pt(16)
    answers_nextrun.font.bold = True
    answers_nextrun.font.name = 'Times New Roman'


    document = Document()
    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run("Cамостоятельная работа")
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'

    nextp = document.add_paragraph()
    nextp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nextrun = nextp.add_run(theme)
    nextrun.font.size = Pt(16)
    nextrun.font.bold = True
    nextrun.font.name = 'Times New Roman'
    points = 'абвгдежзиклмн'
    for i in range(number_of_exercises):
        nextp = document.add_paragraph()
        nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        nextrun = nextp.add_run('№ ' + str(i + 1) + ". Выполните арифметические операции:")
        nextrun.font.size = Pt(14)
        nextrun.font.bold = False
        nextrun.font.name = 'Times New Roman'

        answers_nextp = answers_document.add_paragraph()
        answers_nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        answers_nextrun = answers_nextp.add_run('№ ' + str(i + 1) + ".")
        answers_nextrun.font.size = Pt(14)
        answers_nextrun.font.bold = False
        answers_nextrun.font.name = 'Times New Roman'
        exercises = generate_exercise(6, operations[i % 3], radixs[i % 4])
        pprint(exercises)
        for item, point in zip(exercises[0], points):
            nextp = document.add_paragraph()
            nextp.paragraph_format.first_line_indent = Inches(0.5)
            nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            nextrun = nextp.add_run(point + ") " + item[0])
            nextrun.subscript = False
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'
            subnextrun = nextp.add_run(str(radixs[i % 4 ]))
            subnextrun.subscript = True
            nextrun = nextp.add_run(f" {operations[i % 3]} " + item[1])
            subnextrun = nextp.add_run(str(radixs[i % 4]))
            subnextrun.subscript = True
            nextrun.font.size = Pt(14)
            nextrun.font.bold = False
            nextrun.font.name = 'Times New Roman'

        for item, point in zip(exercises[1], points):
            answers_nextp = answers_document.add_paragraph()
            answers_nextp.paragraph_format.first_line_indent = Inches(0.5)
            answers_nextp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            answers_nextrun = answers_nextp.add_run(point + ") " + item)
            answers_nextrun.subscript = False
            answers_nextrun.font.size = Pt(14)
            answers_nextrun.font.bold = False
            answers_nextrun.font.name = 'Times New Roman'
            answers_subnextrun = answers_nextp.add_run(str(radixs[i % 4 ]))
            answers_subnextrun.subscript = True




    document.save(theme + '.docx')
    convert(theme + '.docx')

    answers_document.save(theme + " ОТВЕТЫ.docx")
    convert(theme + " ОТВЕТЫ.docx")


generate_pdf("Арифметические операции", 6, ['+', '-', "*"], [2, 4, 8, 16])

print(eval("35" + "-" + "43"))